"""
Document generation service for creating legal documents from templates.
"""

from datetime import datetime
from typing import Dict, Any, Optional
from pathlib import Path
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT


class DocumentGenerator:
    """Service for generating legal documents from templates."""

    TEMPLATE_TYPES = {
        "board_resolution": "Board Resolution",
        "meeting_minutes": "Board Meeting Minutes",
        "notice": "Notice of Meeting",
        "consent_action": "Action by Written Consent"
    }

    def __init__(self):
        self.templates_dir = Path(__file__).parent.parent / "templates"
        self.templates_dir.mkdir(exist_ok=True)

    def generate_document(
        self,
        template_type: str,
        data: Dict[str, Any],
        format: str = "docx"
    ) -> bytes:
        """
        Generate a document from a template.

        Args:
            template_type: Type of template (board_resolution, meeting_minutes, etc.)
            data: Dictionary containing template variables
            format: Output format (docx or pdf)

        Returns:
            Document content as bytes
        """
        if template_type not in self.TEMPLATE_TYPES:
            raise ValueError(f"Unknown template type: {template_type}")

        if format not in ["docx", "pdf"]:
            raise ValueError(f"Unsupported format: {format}")

        # Generate document based on template type
        if template_type == "board_resolution":
            content = self._generate_board_resolution(data)
        elif template_type == "meeting_minutes":
            content = self._generate_meeting_minutes(data)
        elif template_type == "notice":
            content = self._generate_notice(data)
        elif template_type == "consent_action":
            content = self._generate_consent_action(data)
        else:
            raise ValueError(f"Template generation not implemented for: {template_type}")

        # Convert to requested format
        if format == "docx":
            return self._generate_docx(content, data)
        else:  # pdf
            return self._generate_pdf(content, data)

    def _generate_board_resolution(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate content for board resolution."""
        return {
            "title": "BOARD RESOLUTION",
            "company": data.get("company", "Atlas Machine and Supply, Inc."),
            "resolution_number": data.get("resolution_number", ""),
            "date": data.get("date", datetime.now().strftime("%B %d, %Y")),
            "sections": [
                {
                    "heading": "RESOLUTION",
                    "content": data.get("resolution_title", "")
                },
                {
                    "heading": "WHEREAS",
                    "content": data.get("whereas_clauses", [])
                },
                {
                    "heading": "NOW, THEREFORE, BE IT RESOLVED",
                    "content": data.get("resolved_clauses", [])
                },
                {
                    "heading": "CERTIFICATION",
                    "content": [
                        f"I hereby certify that the foregoing resolution was duly adopted by the Board of Directors of {data.get('company', 'Atlas Machine and Supply, Inc.')} on {data.get('date', datetime.now().strftime('%B %d, %Y'))}.",
                        "",
                        "_" * 40,
                        data.get("secretary_name", "Secretary"),
                        "Corporate Secretary"
                    ]
                }
            ]
        }

    def _generate_meeting_minutes(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate content for board meeting minutes."""
        return {
            "title": "MINUTES OF BOARD OF DIRECTORS MEETING",
            "company": data.get("company", "Atlas Machine and Supply, Inc."),
            "date": data.get("date", datetime.now().strftime("%B %d, %Y")),
            "time": data.get("time", ""),
            "location": data.get("location", ""),
            "sections": [
                {
                    "heading": "ATTENDANCE",
                    "content": [
                        f"Present: {', '.join(data.get('attendees', []))}",
                        f"Absent: {', '.join(data.get('absent', []))}",
                        f"Also Present: {', '.join(data.get('guests', []))}" if data.get('guests') else ""
                    ]
                },
                {
                    "heading": "CALL TO ORDER",
                    "content": [
                        f"The meeting was called to order at {data.get('time', '')} by {data.get('chair', 'the Chairman')}."
                    ]
                },
                {
                    "heading": "APPROVAL OF MINUTES",
                    "content": data.get("minutes_approval", [
                        "The minutes of the previous meeting were reviewed and approved."
                    ])
                },
                {
                    "heading": "MATTERS DISCUSSED",
                    "content": data.get("matters_discussed", [])
                },
                {
                    "heading": "RESOLUTIONS ADOPTED",
                    "content": data.get("resolutions", [])
                },
                {
                    "heading": "ADJOURNMENT",
                    "content": [
                        f"There being no further business, the meeting was adjourned at {data.get('adjournment_time', '')}."
                    ]
                },
                {
                    "heading": "CERTIFICATION",
                    "content": [
                        "",
                        "_" * 40,
                        data.get("secretary_name", "Secretary"),
                        "Corporate Secretary",
                        "",
                        f"Date: {data.get('date', datetime.now().strftime('%B %d, %Y'))}"
                    ]
                }
            ]
        }

    def _generate_notice(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate content for meeting notice."""
        return {
            "title": "NOTICE OF BOARD OF DIRECTORS MEETING",
            "company": data.get("company", "Atlas Machine and Supply, Inc."),
            "sections": [
                {
                    "heading": "TO",
                    "content": [f"All Directors of {data.get('company', 'Atlas Machine and Supply, Inc.')}"]
                },
                {
                    "heading": "NOTICE IS HEREBY GIVEN",
                    "content": [
                        f"A meeting of the Board of Directors will be held on:",
                        "",
                        f"Date: {data.get('meeting_date', '')}",
                        f"Time: {data.get('meeting_time', '')}",
                        f"Location: {data.get('meeting_location', '')}",
                        "",
                        "The purpose of the meeting is to consider and act upon the following matters:",
                    ] + data.get("agenda_items", [])
                },
                {
                    "heading": "",
                    "content": [
                        "",
                        f"Dated: {data.get('date', datetime.now().strftime('%B %d, %Y'))}",
                        "",
                        "_" * 40,
                        data.get("secretary_name", "Corporate Secretary"),
                        "Secretary"
                    ]
                }
            ]
        }

    def _generate_consent_action(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate content for consent action."""
        return {
            "title": "ACTION BY WRITTEN CONSENT OF THE BOARD OF DIRECTORS",
            "company": data.get("company", "Atlas Machine and Supply, Inc."),
            "date": data.get("date", datetime.now().strftime("%B %d, %Y")),
            "sections": [
                {
                    "heading": "",
                    "content": [
                        f"The undersigned, being all of the Directors of {data.get('company', 'Atlas Machine and Supply, Inc.')}, hereby consent to the following action(s) without a meeting:"
                    ]
                },
                {
                    "heading": "RESOLVED",
                    "content": data.get("resolutions", [])
                },
                {
                    "heading": "SIGNATURES",
                    "content": [
                        "",
                        "This action is effective as of the date signed by all Directors.",
                        ""
                    ] + [
                        f"\n_" + "_" * 40 + f"\n{director}\nDirector\nDate: _____________\n"
                        for director in data.get("directors", [])
                    ]
                }
            ]
        }

    def _generate_docx(self, content: Dict[str, Any], data: Dict[str, Any]) -> bytes:
        """Generate DOCX document from content."""
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Add title
        title = doc.add_heading(content["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add company name
        company = doc.add_heading(content["company"], 2)
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date if present
        if "date" in content:
            date_para = doc.add_paragraph(content["date"])
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add resolution number if present
        if "resolution_number" in content and content["resolution_number"]:
            res_num = doc.add_paragraph(f"Resolution No. {content['resolution_number']}")
            res_num.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Add spacing

        # Add sections
        for section in content["sections"]:
            if section["heading"]:
                doc.add_heading(section["heading"], 3)

            if isinstance(section["content"], list):
                for item in section["content"]:
                    if item:  # Skip empty strings
                        doc.add_paragraph(item)
            else:
                doc.add_paragraph(section["content"])

            doc.add_paragraph()  # Add spacing between sections

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_pdf(self, content: Dict[str, Any], data: Dict[str, Any]) -> bytes:
        """Generate PDF document from content."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # Container for the 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=12
        )
        company_style = ParagraphStyle(
            'CompanyStyle',
            parent=styles['Heading2'],
            fontSize=14,
            alignment=TA_CENTER,
            spaceAfter=12
        )
        heading_style = ParagraphStyle(
            'HeadingStyle',
            parent=styles['Heading3'],
            fontSize=12,
            bold=True,
            spaceAfter=6
        )
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=6
        )

        # Add title
        elements.append(Paragraph(content["title"], title_style))
        elements.append(Spacer(1, 6))

        # Add company name
        elements.append(Paragraph(content["company"], company_style))
        elements.append(Spacer(1, 6))

        # Add date if present
        if "date" in content:
            date_style = ParagraphStyle('DateStyle', parent=body_style, alignment=TA_CENTER)
            elements.append(Paragraph(content["date"], date_style))
            elements.append(Spacer(1, 6))

        # Add resolution number if present
        if "resolution_number" in content and content["resolution_number"]:
            res_style = ParagraphStyle('ResNumStyle', parent=body_style, alignment=TA_CENTER)
            elements.append(Paragraph(f"Resolution No. {content['resolution_number']}", res_style))
            elements.append(Spacer(1, 12))

        elements.append(Spacer(1, 12))

        # Add sections
        for section in content["sections"]:
            if section["heading"]:
                elements.append(Paragraph(section["heading"], heading_style))
                elements.append(Spacer(1, 6))

            if isinstance(section["content"], list):
                for item in section["content"]:
                    if item:  # Skip empty strings
                        elements.append(Paragraph(item, body_style))
                        elements.append(Spacer(1, 3))
            else:
                elements.append(Paragraph(section["content"], body_style))
                elements.append(Spacer(1, 3))

            elements.append(Spacer(1, 12))

        # Build PDF
        doc.build(elements)

        # Get the value of the BytesIO buffer and return
        buffer.seek(0)
        return buffer.getvalue()

    def get_template_fields(self, template_type: str) -> Dict[str, Any]:
        """
        Get the required and optional fields for a template type.

        Returns:
            Dictionary describing fields needed for the template
        """
        templates = {
            "board_resolution": {
                "required": ["resolution_title", "resolved_clauses"],
                "optional": ["company", "resolution_number", "date", "whereas_clauses", "secretary_name"],
                "fields": {
                    "company": {"type": "string", "description": "Company name"},
                    "resolution_number": {"type": "string", "description": "Resolution number/identifier"},
                    "date": {"type": "string", "description": "Date of resolution (MM/DD/YYYY)"},
                    "resolution_title": {"type": "string", "description": "Title of the resolution"},
                    "whereas_clauses": {"type": "array", "description": "List of WHEREAS clauses"},
                    "resolved_clauses": {"type": "array", "description": "List of RESOLVED clauses"},
                    "secretary_name": {"type": "string", "description": "Name of corporate secretary"}
                }
            },
            "meeting_minutes": {
                "required": ["attendees", "matters_discussed"],
                "optional": ["company", "date", "time", "location", "absent", "guests", "chair",
                            "minutes_approval", "resolutions", "adjournment_time", "secretary_name"],
                "fields": {
                    "company": {"type": "string", "description": "Company name"},
                    "date": {"type": "string", "description": "Meeting date"},
                    "time": {"type": "string", "description": "Meeting start time"},
                    "location": {"type": "string", "description": "Meeting location"},
                    "attendees": {"type": "array", "description": "List of attendees"},
                    "absent": {"type": "array", "description": "List of absent members"},
                    "guests": {"type": "array", "description": "List of guests"},
                    "chair": {"type": "string", "description": "Meeting chair/president"},
                    "minutes_approval": {"type": "array", "description": "Minutes approval text"},
                    "matters_discussed": {"type": "array", "description": "List of matters discussed"},
                    "resolutions": {"type": "array", "description": "List of resolutions adopted"},
                    "adjournment_time": {"type": "string", "description": "Adjournment time"},
                    "secretary_name": {"type": "string", "description": "Name of corporate secretary"}
                }
            },
            "notice": {
                "required": ["meeting_date", "meeting_time", "meeting_location", "agenda_items"],
                "optional": ["company", "date", "secretary_name"],
                "fields": {
                    "company": {"type": "string", "description": "Company name"},
                    "date": {"type": "string", "description": "Notice date"},
                    "meeting_date": {"type": "string", "description": "Date of meeting"},
                    "meeting_time": {"type": "string", "description": "Time of meeting"},
                    "meeting_location": {"type": "string", "description": "Location of meeting"},
                    "agenda_items": {"type": "array", "description": "List of agenda items"},
                    "secretary_name": {"type": "string", "description": "Name of corporate secretary"}
                }
            },
            "consent_action": {
                "required": ["resolutions", "directors"],
                "optional": ["company", "date"],
                "fields": {
                    "company": {"type": "string", "description": "Company name"},
                    "date": {"type": "string", "description": "Effective date"},
                    "resolutions": {"type": "array", "description": "List of resolutions"},
                    "directors": {"type": "array", "description": "List of director names for signatures"}
                }
            }
        }

        if template_type not in templates:
            raise ValueError(f"Unknown template type: {template_type}")

        return templates[template_type]
