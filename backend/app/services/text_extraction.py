import os
from typing import Optional, List, Tuple
from pathlib import Path
import tiktoken

# PDF extraction
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

# OCR for scanned PDFs
try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    pytesseract = None
    convert_from_path = None
    Image = None
    OCR_AVAILABLE = False

# Word document extraction
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Excel extraction
try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None


class TextExtractionService:
    """Service for extracting text from various document types"""

    @staticmethod
    def extract_text_from_pdf_with_ocr(file_path: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from PDF using OCR (for scanned/image-based PDFs)

        Returns:
            (extracted_text, error_message)
        """
        if not OCR_AVAILABLE:
            return "", "OCR not available. Install pytesseract and pdf2image."

        try:
            # Convert PDF pages to images
            images = convert_from_path(file_path, dpi=300)

            text_parts = []
            for page_num, image in enumerate(images, start=1):
                try:
                    # Use Tesseract to extract text from image
                    page_text = pytesseract.image_to_string(image, lang='eng')
                    if page_text and page_text.strip():
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                except Exception as e:
                    text_parts.append(f"[Page {page_num}] OCR error: {str(e)}")

            extracted_text = "\n\n".join(text_parts)

            if not extracted_text.strip():
                return "", "OCR completed but no text could be extracted from images"

            return extracted_text, None

        except Exception as e:
            return "", f"OCR extraction error: {str(e)}"

    @staticmethod
    def extract_text_from_pdf(file_path: str, use_ocr_fallback: bool = True) -> Tuple[str, Optional[str]]:
        """
        Extract text from PDF file

        Returns:
            (extracted_text, error_message)
        """
        if not PdfReader:
            return "", "PyPDF2 not installed"

        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                except Exception as e:
                    text_parts.append(f"[Page {page_num}] Error extracting: {str(e)}")

            extracted_text = "\n\n".join(text_parts)

            # If no text was extracted and OCR is available, try OCR
            if not extracted_text.strip() and use_ocr_fallback and OCR_AVAILABLE:
                print(f"No text found in PDF using PyPDF2, attempting OCR for: {file_path}")
                return TextExtractionService.extract_text_from_pdf_with_ocr(file_path)

            if not extracted_text.strip():
                if OCR_AVAILABLE:
                    return "", "No text could be extracted from PDF even with OCR"
                else:
                    return "", "No text could be extracted from PDF. This may be a scanned/image-based PDF that requires OCR (Optical Character Recognition). OCR dependencies are not installed."

            return extracted_text, None

        except Exception as e:
            return "", f"PDF extraction error: {str(e)}"

    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from Word document

        Returns:
            (extracted_text, error_message)
        """
        if not DocxDocument:
            return "", "python-docx not installed"

        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("\n[Table]\n" + "\n".join(table_text))

            extracted_text = "\n\n".join(text_parts)

            if not extracted_text.strip():
                return "", "No text could be extracted from Word document"

            return extracted_text, None

        except Exception as e:
            return "", f"DOCX extraction error: {str(e)}"

    @staticmethod
    def extract_text_from_xlsx(file_path: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from Excel spreadsheet

        Returns:
            (extracted_text, error_message)
        """
        if not load_workbook:
            return "", "openpyxl not installed"

        try:
            workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"[Sheet: {sheet_name}]"]

                # Extract rows
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [str(cell) for cell in row if cell is not None]
                    if row_values:
                        sheet_text.append(" | ".join(row_values))

                if len(sheet_text) > 1:  # Has content beyond header
                    text_parts.append("\n".join(sheet_text))

            extracted_text = "\n\n".join(text_parts)

            if not extracted_text.strip():
                return "", "No text could be extracted from Excel file"

            return extracted_text, None

        except Exception as e:
            return "", f"Excel extraction error: {str(e)}"

    @staticmethod
    def extract_text(file_path: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from document based on file extension

        Returns:
            (extracted_text, error_message)
        """
        if not os.path.exists(file_path):
            return "", "File not found"

        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return TextExtractionService.extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return TextExtractionService.extract_text_from_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return TextExtractionService.extract_text_from_xlsx(file_path)
        else:
            return "", f"Unsupported file type: {ext}"

    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        encoding_name: str = "cl100k_base"
    ) -> List[str]:
        """
        Split text into chunks based on token count

        Args:
            text: Text to chunk
            chunk_size: Target tokens per chunk
            chunk_overlap: Overlap between chunks in tokens
            encoding_name: Tiktoken encoding (cl100k_base for Claude)

        Returns:
            List of text chunks
        """
        try:
            encoding = tiktoken.get_encoding(encoding_name)
        except Exception:
            # Fallback to character-based chunking
            return TextExtractionService._chunk_by_characters(
                text, chunk_size * 4, chunk_overlap * 4
            )

        # Tokenize the entire text
        tokens = encoding.encode(text)

        if len(tokens) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(tokens):
            # Get chunk
            end = start + chunk_size
            chunk_tokens = tokens[start:end]

            # Decode back to text
            chunk_text = encoding.decode(chunk_tokens)
            chunks.append(chunk_text)

            # Move start forward (with overlap)
            start = end - chunk_overlap

            # Prevent infinite loop
            if start >= len(tokens):
                break

        return chunks

    @staticmethod
    def _chunk_by_characters(
        text: str,
        chunk_size: int = 2000,
        chunk_overlap: int = 200
    ) -> List[str]:
        """
        Fallback chunking by character count

        Args:
            text: Text to chunk
            chunk_size: Characters per chunk
            chunk_overlap: Overlap between chunks

        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]

            # Try to break at paragraph boundary
            if end < len(text):
                last_break = max(
                    chunk.rfind('\n\n'),
                    chunk.rfind('\n'),
                    chunk.rfind('. '),
                    chunk.rfind('! '),
                    chunk.rfind('? ')
                )
                if last_break > chunk_size // 2:  # Only break if in second half
                    chunk = chunk[:last_break + 1]
                    end = start + last_break + 1

            chunks.append(chunk)
            start = end - chunk_overlap

            if start >= len(text):
                break

        return chunks

    @staticmethod
    def count_tokens(text: str, encoding_name: str = "cl100k_base") -> int:
        """
        Count tokens in text

        Args:
            text: Text to count
            encoding_name: Tiktoken encoding

        Returns:
            Token count
        """
        try:
            encoding = tiktoken.get_encoding(encoding_name)
            return len(encoding.encode(text))
        except Exception:
            # Fallback: estimate ~4 characters per token
            return len(text) // 4
